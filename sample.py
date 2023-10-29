import language_tool_python
from docx import Document

# Load the LanguageTool API
tool = language_tool_python.LanguageTool('en-US')

# Load the Word document
document = Document('sample.docx')

# Iterate over paragraphs in the document
for paragraph in document.paragraphs:
    text = paragraph.text
    # Perform grammar and spell checking
    matches = tool.check(text)
    
    # Iterate over the matches and highlight the errors
    for match in matches:
        start_index = match.offset
        end_index = start_index + match.errorLength
        # Highlight the error using Word's built-in style
        paragraph.add_run(text[start_index:end_index]).bold = True

# Save the updated document
document.save('sample_with_errors.docx')
